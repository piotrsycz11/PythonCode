import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import requests
import json
import docx

# API key do Eden AI
API_KEY = 'Twoj_klucz_API'

full_text = ""
sciezka_pliku = None

# Funkcja wyboru pliku docx
def wybierz_plik_docx():
    global sciezka_pliku
    global full_text
    sciezka_pliku = filedialog.askopenfilename(filetypes=[("Pliki Word", "*.docx")])
    if sciezka_pliku:
        doc = docx.Document(sciezka_pliku)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        full_text = '\n'.join(full_text)

# Funkcja tłumaczenia tekstu
def tlumacz_tekst():
    # Zmienne dla języków mogą być zmienione w zależności od wyboru użytkownika
    source_lang = var.get().split('-')[0]
    target_lang = var.get().split('-')[1]
    global full_text

    if not full_text:  # Sprawdzenie czy tekst nie jest pusty
        messagebox.showerror("Błąd", "Tekst do tłumaczenia jest pusty.")
        return
    # Kod proponowany przez API
    headers = {"Authorization": f"Bearer {API_KEY}"}
    url = "https://api.edenai.run/v2/translation/automatic_translation"
    payload = {
        "providers": "google",
        "source_language": source_lang,
        "target_language": target_lang,
        "text": full_text,
    }

    response = requests.post(url, json=payload, headers=headers)
    if response.status_code == 200:
        result = json.loads(response.text)
        zapisz_do_pliku_docx(result)
    else:
        messagebox.showerror("Błąd", f"Nie udało się przetłumaczyć tekstu: {response.text}")

# Funkcja zapisu przetłumaczonego tekstu do pliku
def zapisz_do_pliku_docx(translated_text):
    new_doc = docx.Document()
    new_doc.add_paragraph(translated_text['google']['text'])
    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("Pliki Word", "*.docx"), ("Wszystkie pliki", "*.*")])
    if file_path:  # Jeśli ścieżka jest niepusta
        new_doc.save(file_path)
        messagebox.showinfo("Sukces", "Przetłumaczony dokument został zapisany.")

# Inicjalizacja 
root = tk.Tk()
root.title("File Translator")
root.geometry("300x150")

# Ramka główna
main_frame = ttk.Frame(root)
main_frame.pack(fill='both', expand=False)

# Ramka z wyborem języków
language_frame = ttk.Frame(main_frame)
language_frame.pack(pady=10)

# Ramka z przyciskami
button_frame = ttk.Frame(main_frame)
button_frame.pack(pady=10)

# Wybór języków
languages_label = ttk.Label(language_frame, text="Wybierz języki:")
languages_label.pack(side='left', padx=10)

var = tk.StringVar(value="pl-en")
pl_en_radio = ttk.Radiobutton(language_frame, text="PL -> ENG", variable=var, value="pl-en")
en_pl_radio = ttk.Radiobutton(language_frame, text="ENG -> PL", variable=var, value="en-pl")

pl_en_radio.pack(side='left', padx=10)
en_pl_radio.pack(side='left', padx=10)

# Przyciski
choose_button = ttk.Button(button_frame, text="Wybierz plik", command=wybierz_plik_docx)
translate_button = ttk.Button(button_frame, text="Tłumacz", command=tlumacz_tekst)

choose_button.pack(side='left', padx=10)
translate_button.pack(side='left', padx=10)

root.mainloop()